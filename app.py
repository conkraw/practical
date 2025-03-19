import streamlit as st
import pandas as pd
from docx import Document
import io
import zipfile
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_document(student_name, code, exam_number):
    """
    Generate a Word document for a given exam number.
    """
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Add heading for the specific practical examination
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f'Pediatric Clerkship Practical Examination #{exam_number}')
    run.bold = True
    run.font.all_caps = True  # This will display the text in all uppercase
    run.underline = True  
        
    # Add a paragraph with the student's legal name (with the name in bold)
    paragraph = doc.add_paragraph()
    paragraph.add_run("Student Name: ")
    run = paragraph.add_run(str(student_name))
    run.bold = True

    # Add access instructions heading
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Practical Examination Access Instructions")
    run.bold = True
    run.underline = True
    
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # ---------------------------
    # Exam Access Instructions
    # ---------------------------
    doc.add_paragraph("Exam Access Instructions", style='Heading2')
    
    # Step 1: Website URL
    doc.add_paragraph("1. Visit the exam website: https://redcap.ctsi.psu.edu/surveys")
    
    # Step 2: Code Entry (with bold code)
    paragraph = doc.add_paragraph()
    paragraph.add_run("2. Enter the exam code: ")
    code_run = paragraph.add_run(str(code))
    code_run.bold = True
    
    # ---------------------------
    # Test Taking Instructions
    # ---------------------------
    doc.add_paragraph("Test Taking Instructions", style='Heading2')
    doc.add_paragraph("3. You can navigate backward and forward through the exam questions.")
    doc.add_paragraph("4. If your computer fails, your last response will be saved automatically. Resume by clicking 'Next'.")
    doc.add_paragraph("5. Once the exam is submitted, it cannot be reopened.")
    doc.add_paragraph("6. You have 1 hour to complete the exam; a proctor will monitor the timer.")
    doc.add_paragraph("7. Use the erasable noteboard provided. All noteboards must be returned at the end of the exam.")
    doc.add_paragraph("8. The calculator app on your computer is permitted; phones are not allowed.")
    doc.add_paragraph("9. Screenshots or any form of screen capture of the exam content are strictly prohibited.")
    
    # ---------------------------
    # Exam Overview / Introduction
    # ---------------------------
    intro_text = (
        "Welcome to the Pediatric Clerkship Practical Examination!\n\n"
        "This exam offers you the opportunity to demonstrate your skills in biomedical knowledge, clinical reasoning, "
        "and systems-based thinking as you work through simulated cases involving undifferentiated pediatric patients.\n\n"
        "You will be presented with a series of questions. All fields must be completed to proceed. The purpose of this exam is "
        "to assess your ability to synthesize data, interpret findings, and apply critical thinking in clinical decision-making.\n\n"
        "All necessary resources, including immunization schedules and pharmacologic references, will be available during the exam. "
        "Please note that any notes taken cannot be removed from the exam room."
    )
    intro_paragraph = doc.add_paragraph(intro_text)
    intro_paragraph.paragraph_format.space_after = Pt(0)
    
    # ---------------------------
    # Learning Objectives
    # ---------------------------
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run_title = paragraph.add_run("Learning Objectives")
    run_title.bold = True
    run_title.add_break()  # Line break so that objectives start on the next line
    
    paragraph.add_run(
        "• Demonstrate comprehensive medical history taking for patients of all ages.\n"
        "• Formulate assessments, differential diagnoses, and management plans for common pediatric complaints.\n"
        "• Apply knowledge of growth and development to create individualized pediatric care strategies.\n"
        "• Evaluate social determinants of health and their impact on pediatric outcomes.\n\n"
        "Good luck, and let this exam be an opportunity to showcase your professionalism and clinical expertise."
    )


st.title("Pediatric Clerkship Practical Examination Document Generator")

# File uploader widget
uploaded_file = st.file_uploader("Upload your CSV file", type="csv")

if uploaded_file is not None:
    # Read the CSV file into a DataFrame
    df = pd.read_csv(uploaded_file)
    st.write("Data Preview:", df.head())
    
    # Create in-memory ZIP files for each exam set
    zip_buffer_p1 = io.BytesIO()
    zip_buffer_p2 = io.BytesIO()
    
    # Open both ZIP files for writing
    with zipfile.ZipFile(zip_buffer_p1, "a", zipfile.ZIP_DEFLATED, False) as zip_file_p1, \
         zipfile.ZipFile(zip_buffer_p2, "a", zipfile.ZIP_DEFLATED, False) as zip_file_p2:
        # Process each row in the dataset
        for index, row in df.iterrows():
            student_name = row['legal_name']
            code_p1 = row['code_p1']
            code_p2 = row['code_p2']
            
            # Generate document for Practical Examination #1
            doc1 = generate_document(student_name, code_p1, exam_number=1)
            doc1_io = io.BytesIO()
            doc1.save(doc1_io)
            doc1_io.seek(0)
            filename_p1 = f"{student_name}_p1.docx"
            zip_file_p1.writestr(filename_p1, doc1_io.read())
            
            # Generate document for Practical Examination #2
            doc2 = generate_document(student_name, code_p2, exam_number=2)
            doc2_io = io.BytesIO()
            doc2.save(doc2_io)
            doc2_io.seek(0)
            filename_p2 = f"{student_name}_p2.docx"
            zip_file_p2.writestr(filename_p2, doc2_io.read())
    
    # Reset buffers for download
    zip_buffer_p1.seek(0)
    zip_buffer_p2.seek(0)
    
    # Provide download buttons for both ZIP files
    st.download_button(
        label="Download All Exam #1 Documents as ZIP",
        data=zip_buffer_p1,
        file_name="generated_documents_p1.zip",
        mime="application/zip"
    )
    
    st.download_button(
        label="Download All Exam #2 Documents as ZIP",
        data=zip_buffer_p2,
        file_name="generated_documents_p2.zip",
        mime="application/zip"
    )
